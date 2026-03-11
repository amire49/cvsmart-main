from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import logging
import concurrent.futures
import pdfplumber
from docx import Document
import google.generativeai as genai
import os
import re
import io
import json
import requests
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

app = Flask(__name__)
CORS(app)

# Initialize rate limiter
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

limiter = Limiter(
    key_func=get_remote_address,  # renamed argument
    default_limits=["100 per hour"]
)
limiter.init_app(app)

# Configure logging
logging.basicConfig(level=logging.INFO)

# Configure Gemini API using environment variable
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("No GOOGLE_API_KEY found in environment variables")
genai.configure(api_key=GOOGLE_API_KEY)

ANALYSIS_SYSTEM = """You are a senior recruiter and career coach. You ONLY output valid JSON. No markdown, no code fences, no explanation outside the JSON.

Given a resume and job description, return a single JSON object with these exact keys:

{
  "overallScore": <integer 0-100>,
  "technicalScore": <integer 0-100>,
  "technicalDetails": ["<concise finding>", "<concise finding>", "<concise finding>"],
  "experienceScore": <integer 0-100>,
  "experienceDetails": ["<concise finding>", "<concise finding>", "<concise finding>"],
  "projectScore": <integer 0-100>,
  "projectDetails": ["<concise finding>", "<concise finding>", "<concise finding>"],
  "communicationScore": <integer 0-100>,
  "communicationDetails": ["<concise finding>", "<concise finding>", "<concise finding>"],
  "strengths": ["<key strength 1>", "<key strength 2>", "<key strength 3>"],
  "gaps": ["<key gap 1>", "<key gap 2>", "<key gap 3>"],
  "recommendation": "<2-3 sentence summary of overall fit and top advice>",
  "analysis": "<full detailed analysis as a single string with markdown formatting>"
}

Rules for each field:
- overallScore: genuine 0-100 based on skills match, experience relevance, education, projects. Use the full range.
- technicalScore: how well the candidate's technical skills match the JD requirements (0-100).
- technicalDetails: exactly 3 concise bullets about matched skills, partial matches, and missing skills.
- experienceScore: how relevant the candidate's work experience is (0-100).
- experienceDetails: exactly 3 concise bullets about high-impact experience, experience gaps, years of experience fit.
- projectScore: how well projects/portfolio demonstrate required capabilities (0-100).
- projectDetails: exactly 3 concise bullets about relevant projects, missing project evidence, portfolio quality.
- communicationScore: soft skills, teamwork, communication, cultural fit assessment (0-100).
- communicationDetails: exactly 3 concise bullets about communication, leadership, collaboration indicators.
- strengths: exactly 3 key strengths of this candidate for this role.
- gaps: exactly 3 key gaps or missing qualifications.
- recommendation: 2-3 sentences of professional advice summarizing the match and top actions.
- analysis: a thorough markdown-formatted analysis (use headings, bullets, bold) covering:
  1. Overall Match Score with percentage
  2. Key Skills Match (matched, partial, missing)
  3. Experience Relevance (high-impact, gaps)
  4. Top 3-5 Improvement Tips (specific, actionable)

Keep each detail string under 100 characters. The analysis string can be longer (500-1500 words).
Do NOT invent information not in the resume or job description."""

model = genai.GenerativeModel(model_name='gemini-2.5-flash')

model_analysis = genai.GenerativeModel(
    model_name='gemini-2.5-flash',
    system_instruction=ANALYSIS_SYSTEM,
)


def build_structured_from_json(data):
    """Convert the JSON response from Gemini into the frontend-expected structure."""
    if not isinstance(data, dict):
        return None

    def _int(val, default=50):
        try:
            return max(0, min(100, int(val)))
        except (TypeError, ValueError):
            return default

    def _list(val, fallback="See full analysis"):
        if isinstance(val, list) and val:
            return [str(s).strip()[:100] for s in val if str(s).strip()][:4] or [fallback]
        return [fallback]

    def _str(val, default=""):
        return str(val).strip() if val else default

    overall = _int(data.get("overallScore"), 50)

    if overall >= 75:
        verdict, verdict_color = "Strong Match", "#00e5a0"
    elif overall >= 50:
        verdict, verdict_color = "Conditional Match", "#f4a261"
    else:
        verdict, verdict_color = "Weak Match", "#ff4d6d"

    return {
        "overallScore": overall,
        "verdict": verdict,
        "verdictColor": verdict_color,
        "sections": [
            {
                "id": "technical",
                "label": "Technical Alignment",
                "score": _int(data.get("technicalScore"), overall),
                "color": "#00e5a0",
                "icon": "⚡",
                "details": _list(data.get("technicalDetails")),
            },
            {
                "id": "experience",
                "label": "Experience Match",
                "score": _int(data.get("experienceScore"), overall),
                "color": "#ff4d6d",
                "icon": "📅",
                "details": _list(data.get("experienceDetails")),
            },
            {
                "id": "projects",
                "label": "Project Relevance",
                "score": _int(data.get("projectScore"), overall),
                "color": "#00b4d8",
                "icon": "🏗️",
                "details": _list(data.get("projectDetails")),
            },
            {
                "id": "softskills",
                "label": "Communication & Fit",
                "score": _int(data.get("communicationScore"), overall),
                "color": "#f4a261",
                "icon": "💬",
                "details": _list(data.get("communicationDetails")),
            },
        ],
        "strengths": _list(data.get("strengths"), "See full analysis for strengths")[:3],
        "gaps": _list(data.get("gaps"), "See full analysis for gaps")[:3],
        "recommendation": _str(data.get("recommendation"),
                               f"Overall match score: {overall}%. See full analysis for details.")[:300],
    }

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_keywords_from_resume(resume_text, job_description="", max_keywords=12):
    """Use Gemini to extract job-relevant keywords and suggested job titles from resume (and optional job description)."""
    if not resume_text or len(resume_text.strip()) < 50:
        words = re.findall(r'[a-zA-Z]{3,}', resume_text[:500])
        return list(dict.fromkeys(words))[:max_keywords] if words else ["software", "developer"], []

    jd_context = ""
    if job_description and len(job_description.strip()) > 20:
        jd_context = f"\n\nTarget job description (use to align keywords):\n{job_description[:2000]}"

    try:
        prompt = f"""From this resume text, extract:
1. Exactly {max_keywords} job-relevant skills, technologies, or role keywords (e.g. Python, project management, AWS). One per line.
2. Then on a new line write "JOB_TITLES:" followed by 2-3 suggested job titles the candidate could search for (e.g. Software Engineer, Full Stack Developer). Comma-separated.
Output format: keywords first (one per line), then "JOB_TITLES: title1, title2, title3". Only the keywords and job titles, nothing else."""
        response = model.generate_content(
            prompt + "\n\nResume:\n" + resume_text[:4000] + jd_context,
            generation_config=genai.GenerationConfig(temperature=0.1, max_output_tokens=300),
        )
        text = (response.text or "").strip()
        keywords = []
        job_titles = []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.upper().startswith("JOB_TITLES:"):
                titles_str = line.split(":", 1)[-1].strip()
                job_titles = [t.strip() for t in titles_str.split(",") if t.strip() and len(t.strip()) < 60]
                break
            if len(line) < 50 and line not in keywords:
                keywords.append(line)
        keywords = keywords[:max_keywords]
        return keywords if keywords else ["software", "developer"], job_titles[:3]
    except Exception:
        words = re.findall(r'[a-zA-Z]{3,}', resume_text[:500])
        return list(dict.fromkeys(words))[:max_keywords] if words else ["software", "developer"], []


def _parse_resume_heuristic(resume_text, job_description=""):
    """Best-effort parse of raw resume text into structured sections."""
    lines = (resume_text or "").strip().splitlines()
    name = ""
    email = ""
    phone = ""
    location = ""

    for line in lines[:15]:
        stripped = line.strip()
        if not stripped:
            continue
        if not name and len(stripped) < 60 and not re.search(r'[@\d]', stripped):
            name = stripped
            continue
        em = re.search(r'[\w.+-]+@[\w.-]+\.\w+', stripped)
        if em and not email:
            email = em.group(0)
        ph = re.search(r'\+?[\d\s\-().]{7,20}', stripped)
        if ph and not phone:
            phone = ph.group(0).strip()
        if any(kw in stripped.lower() for kw in ['addis', 'ababa', 'nairobi', 'new york', 'london', 'remote', 'ethiopia']):
            if not location:
                location = stripped

    section_map = {}
    current_section = "other"
    section_keywords = {
        "summary": ["summary", "professional summary", "profile", "objective", "about"],
        "experience": ["experience", "work experience", "employment", "work history"],
        "education": ["education", "academic", "qualification"],
        "skills": ["skills", "technical skills", "competencies", "technologies"],
        "projects": ["projects", "portfolio", "key projects"],
        "hobbies": ["hobbies", "interests", "activities"],
    }

    for line in lines:
        stripped = line.strip()
        low = stripped.lower().rstrip(':')
        matched_section = None
        for sec_key, keywords in section_keywords.items():
            if low in keywords or any(low == kw for kw in keywords):
                matched_section = sec_key
                break
        if matched_section:
            current_section = matched_section
            section_map.setdefault(current_section, [])
            continue
        if stripped:
            section_map.setdefault(current_section, []).append(stripped)

    skills_list = []
    for s_line in section_map.get("skills", []):
        for part in re.split(r'[,|;]', s_line):
            part = part.strip().strip('-').strip()
            if part and len(part) < 50:
                skills_list.append(part)

    experience_list = []
    exp_lines = section_map.get("experience", [])
    current_exp = None
    for el in exp_lines:
        if len(el) < 80 and not el.startswith('-') and not el.startswith('*'):
            if current_exp:
                experience_list.append(current_exp)
            current_exp = {"role": el, "company": "", "location": "", "dates": "", "bullets": []}
        elif current_exp:
            current_exp["bullets"].append(el.lstrip('-').lstrip('*').strip())
        else:
            current_exp = {"role": "", "company": "", "location": "", "dates": "", "bullets": [el.lstrip('-').lstrip('*').strip()]}
    if current_exp:
        experience_list.append(current_exp)

    education_list = []
    for ed_line in section_map.get("education", []):
        education_list.append({"degree": ed_line, "school": "", "year": ""})

    summary_text = " ".join(section_map.get("summary", []))[:500]
    if not summary_text:
        jd_title = (job_description or "").splitlines()[0].strip()[:100] if job_description else ""
        summary_text = f"Experienced professional applying for {jd_title}." if jd_title else ""

    projects_list = []
    for pl in section_map.get("projects", []):
        projects_list.append({"title": pl, "description": ""})

    hobbies = section_map.get("hobbies", [])

    return {
        "name": name,
        "title": "",
        "contact": {"email": email, "phone": phone, "location": location, "website": ""},
        "summary": summary_text,
        "skills": skills_list[:20],
        "experience": experience_list,
        "education": education_list,
        "projects": projects_list,
        "hobbies": hobbies,
    }


def generate_structured_cv_sections(resume_text, job_description, analysis_text=""):
    """Use Gemini to turn resume + job description into structured JSON CV sections."""
    prompt_parts = [
        "Resume text:\n", resume_text[:6000] or "",
        "\n\nJob description:\n", job_description[:4000] or "",
    ]
    if analysis_text:
        prompt_parts.append("\n\nFeedback to incorporate:\n")
        prompt_parts.append(analysis_text[:3000])
    prompt_parts.append(
        "\n\nReturn ONLY a valid JSON object with keys: name, title, contact, summary, "
        "skills, experience, education, projects, hobbies. No markdown fences."
    )
    full_prompt = "".join(prompt_parts)

    data = {}
    try:
        response = model_cv_json.generate_content(
            full_prompt,
            generation_config=genai.GenerationConfig(
                temperature=0.15,
                max_output_tokens=4000,
            ),
        )
        text = (response.text or "").strip()
        if "```" in text:
            parts = re.split(r"```(?:\w*)\s*\n?", text, maxsplit=1)
            if len(parts) > 1:
                text = parts[1]
            text = re.sub(r"\s*```\s*$", "", text)
        text = text.strip()
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            text = text[start : end + 1]
        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            text_fixed = re.sub(r",\s*([}\]])", r"\1", text)
            try:
                data = json.loads(text_fixed)
            except json.JSONDecodeError:
                logging.warning("CV JSON parse failed; raw (first 400): %s", text[:400])
                data = {}
    except Exception as e:
        logging.exception("generate_structured_cv_sections Gemini error: %s", e)

    if not isinstance(data, dict) or not data or not data.get("name"):
        logging.info("Gemini JSON unusable, falling back to heuristic parsing")
        data = _parse_resume_heuristic(resume_text, job_description)

    def _s(value, default=""):
        return value if isinstance(value, str) else default

    def _l(value):
        return value if isinstance(value, list) else []

    contact = data.get("contact") or {}
    if not isinstance(contact, dict):
        contact = {}

    normalized = {
        "name": _s(data.get("name")),
        "title": _s(data.get("title")),
        "contact": {
            "email": _s(contact.get("email")),
            "phone": _s(contact.get("phone")),
            "location": _s(contact.get("location")),
            "website": _s(contact.get("website")),
        },
        "summary": _s(data.get("summary")),
        "skills": [str(s).strip() for s in _l(data.get("skills")) if str(s).strip()],
        "experience": [],
        "education": [],
        "projects": [],
        "hobbies": [str(h).strip() for h in _l(data.get("hobbies")) if str(h).strip()],
    }

    for exp in _l(data.get("experience")):
        if not isinstance(exp, dict):
            continue
        normalized["experience"].append({
            "role": _s(exp.get("role")),
            "company": _s(exp.get("company")),
            "location": _s(exp.get("location")),
            "dates": _s(exp.get("dates")),
            "bullets": [str(b).strip() for b in _l(exp.get("bullets")) if str(b).strip()],
        })

    for edu in _l(data.get("education")):
        if not isinstance(edu, dict):
            continue
        normalized["education"].append({
            "degree": _s(edu.get("degree")),
            "school": _s(edu.get("school")),
            "year": _s(edu.get("year")),
        })

    for proj in _l(data.get("projects")):
        if not isinstance(proj, dict):
            continue
        normalized["projects"].append({
            "title": _s(proj.get("title")),
            "description": _s(proj.get("description")),
        })

    return normalized


def fetch_jobs_from_adzuna(keywords, country="gb", per_page=10):
    """Fetch job listings from Adzuna API."""
    app_id = os.getenv("ADZUNA_APP_ID")
    app_key = os.getenv("ADZUNA_APP_KEY")
    if not app_id or not app_key:
        logging.warning("ADZUNA_APP_ID or ADZUNA_APP_KEY not set; returning empty job list")
        return []
    query = " ".join(keywords[:5])
    url = f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
    params = {"app_id": app_id, "app_key": app_key, "what": query, "results_per_page": per_page}
    try:
        resp = requests.get(url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        results = data.get("results") or []
        jobs = []
        for r in results:
            jobs.append({
                "title": r.get("title", ""),
                "company": r.get("company", {}).get("display_name", "") if isinstance(r.get("company"), dict) else str(r.get("company", "")),
                "location": r.get("location", {}).get("display_name", "") if isinstance(r.get("location"), dict) else str(r.get("location", "")),
                "link": r.get("redirect_url", ""),
                "snippet": (r.get("description", "") or "")[:200],
                "source": "adzuna",
            })
        return jobs
    except Exception as e:
        logging.exception("Adzuna API error: %s", e)
        return []


def _normalize_job(job):
    """Normalize job to common shape with source."""
    return {
        "title": (job.get("title") or "").strip(),
        "company": (job.get("company") or "").strip(),
        "location": (job.get("location") or "").strip() if isinstance(job.get("location"), str) else "",
        "link": (job.get("link") or "").strip(),
        "snippet": (job.get("snippet") or "")[:200],
        "source": job.get("source", "unknown"),
    }


def merge_and_deduplicate(job_lists, max_jobs=25):
    """Merge job lists from multiple sources, deduplicate by title+company, cap total."""
    seen = set()
    merged = []
    for lst in job_lists:
        for job in lst:
            n = _normalize_job(job)
            if not n["title"] or not n["link"]:
                continue
            key = (n["title"].lower()[:50], n["company"].lower()[:50])
            if key in seen:
                continue
            seen.add(key)
            merged.append(n)
            if len(merged) >= max_jobs:
                return merged
    return merged


def rank_jobs_by_relevance(jobs, resume_text, max_to_rank=15):
    """Use Gemini to score job relevance (1-10) and sort by score."""
    if not jobs or not resume_text or len(resume_text) < 50:
        return jobs
    to_rank = jobs[:max_to_rank]
    try:
        job_list = "\n".join(
            f"{i+1}. {j['title']} at {j['company']}: {j['snippet'][:80]}..."
            for i, j in enumerate(to_rank)
        )
        prompt = f"""Given this resume summary, score each job's relevance from 1-10. Resume: {resume_text[:6000]}\n\nJobs:\n{job_list}\n\nRespond with exactly {len(to_rank)} numbers, one per line, in order (e.g. 8, 7, 9, ...). Only the numbers."""
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(temperature=0.1, max_output_tokens=100),
        )
        text = (response.text or "").strip()
        scores = []
        for line in text.split("\n"):
            for part in re.split(r"[\s,]+", line.strip()):
                if part.isdigit():
                    scores.append(int(part))
                    if len(scores) >= len(to_rank):
                        break
            if len(scores) >= len(to_rank):
                break
        if len(scores) >= len(to_rank):
            ranked = sorted(zip(to_rank, scores), key=lambda x: -x[1])
            return [j for j, _ in ranked] + jobs[max_to_rank:]
    except Exception as e:
        logging.warning("Relevance ranking failed: %s", e)
    return jobs


def fetch_jobs_from_jsearch(keywords, location="", per_page=10):
    """Fetch job listings from JSearch API (RapidAPI) - aggregates Indeed, LinkedIn, etc."""
    api_key = os.getenv("RAPIDAPI_KEY")
    if not api_key:
        logging.debug("RAPIDAPI_KEY not set; skipping JSearch")
        return []
    query = " ".join(keywords[:5])
    url = "https://jsearch.p.rapidapi.com/search"
    params = {"query": query, "page": "1", "num_pages": "1"}
    if location:
        params["query"] = f"{query} {location}"
    headers = {
        "X-RapidAPI-Key": api_key,
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
    }
    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        results = data.get("data") or []
        jobs = []
        for r in results[:per_page]:
            loc = r.get("job_city") or r.get("job_country") or ""
            if r.get("job_is_remote"):
                loc = "Remote" if not loc else f"{loc} (Remote)"
            jobs.append({
                "title": r.get("job_title") or r.get("title", ""),
                "company": r.get("employer_name") or r.get("company_name", ""),
                "location": loc,
                "link": r.get("job_apply_link") or r.get("job_google_link") or r.get("url", ""),
                "snippet": (r.get("job_description") or r.get("description") or "")[:200],
                "source": "jsearch",
            })
        return jobs
    except Exception as e:
        logging.exception("JSearch API error: %s", e)
        return []


def fetch_jobs_from_arbeitnow(keywords, location="", per_page=10):
    """Fetch job listings from Arbeitnow API (no auth required)."""
    query = " ".join(keywords[:4])
    url = "https://arbeitnow.com/api/job-board-api"
    params = {"search": query}
    if location:
        params["location"] = location
    try:
        resp = requests.get(url, params=params, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        results = data.get("data") or []
        jobs = []
        for r in results[:per_page]:
            slug = r.get("slug", "")
            link = f"https://arbeitnow.com/view/job/{slug}" if slug else r.get("url", "")
            jobs.append({
                "title": r.get("title", ""),
                "company": r.get("company_name", ""),
                "location": (r.get("location") or {}).get("name", "") if isinstance(r.get("location"), dict) else str(r.get("location") or ""),
                "link": link,
                "snippet": re.sub(r"<[^>]+>", "", (r.get("description") or "")[:200]),
                "source": "arbeitnow",
            })
        return jobs
    except Exception as e:
        logging.exception("Arbeitnow API error: %s", e)
        return []


@app.before_request
def require_api_key():
    expected = os.getenv('AUTHORIZED_API_KEY') or ''
    expected = expected.strip()
    if not expected:
        return  # No key configured: allow all requests (e.g. local dev)
    api_key = (request.headers.get('X-API-KEY') or '').strip()
    if api_key != expected:
        return jsonify({'error': 'Unauthorized access'}), 403

@app.before_request
def log_request():
    logging.info(f"Request from {get_remote_address()} to {request.path}")

@app.route('/analyze', methods=['POST'])
@limiter.limit("10 per minute")
def analyze_resume():
    if 'resume' not in request.files or not request.files['resume']:
        return jsonify({'error': 'No resume file uploaded'}), 400
    
    job_description = request.form.get('jobDescription', '')
    if not job_description or len(job_description) < 10:
        return jsonify({'error': 'Invalid job description'}), 400
    
    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Extract text based on file type
            if filename.endswith('.pdf'):
                resume_text = extract_text_from_pdf(file_path)
            else:  # docx
                resume_text = extract_text_from_docx(file_path)
            
            # Clean up the uploaded file
            os.remove(file_path)
            
            prompt = f"Resume:\n{resume_text[:8000]}\n\nJob Description:\n{job_description[:4000]}"

            response = model_analysis.generate_content(
                prompt,
                generation_config=genai.GenerationConfig(
                    temperature=0.05,
                    max_output_tokens=8000,
                    response_mime_type="application/json",
                ),
            )
            text = (response.text or "").strip()

            data = None
            if text:
                if "```" in text:
                    parts = re.split(r"```(?:\w*)\s*\n?", text, maxsplit=1)
                    if len(parts) > 1:
                        text = parts[1]
                    text = re.sub(r"\s*```\s*$", "", text).strip()
                start = text.find("{")
                end = text.rfind("}")
                if start != -1 and end != -1 and end > start:
                    text = text[start : end + 1]
                text_fixed = re.sub(r",\s*([}\]])", r"\1", text)
                try:
                    data = json.loads(text_fixed)
                except json.JSONDecodeError:
                    logging.warning("analyze: JSON parse failed, raw (500 chars): %s", text[:500])

            if isinstance(data, dict):
                analysis = data.get("analysis") or ""
                structured = build_structured_from_json(data)
            else:
                analysis = response.text or ""
                structured = None

            return jsonify({'analysis': analysis, 'structured': structured})
            
        except Exception as e:
            logging.exception("analyze_resume Gemini error: %s", e)
            msg = str(e)
            lower = msg.lower()
            if "429" in msg or "quota" in lower or "rate limit" in lower:
                return jsonify({
                    'error': 'AI quota has been exceeded for now. Please wait a minute and try again.'
                }), 503
            return jsonify({
                'error': 'AI analysis is temporarily unavailable. Please try again in a moment.'
            }), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/jobs/recommend', methods=['POST'])
@limiter.limit("20 per minute")
def jobs_recommend():
    """Recommend jobs based on resume (file or text), aggregating from multiple APIs."""
    resume_text = request.form.get("resumeText", "").strip()
    job_description = request.form.get("jobDescription", "").strip()
    file = request.files.get("resume") if "resume" in request.files else None

    if not resume_text and (not file or not file.filename):
        return jsonify({"error": "Provide resume file or resumeText"}), 400

    if file and file.filename and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        try:
            file.save(file_path)
            if filename.endswith(".pdf"):
                resume_text = extract_text_from_pdf(file_path)
            else:
                resume_text = extract_text_from_docx(file_path)
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)

    if not resume_text or len(resume_text) < 20:
        return jsonify({"error": "Could not extract enough text from resume"}), 400

    keywords, job_titles = extract_keywords_from_resume(resume_text, job_description)
    search_keywords = keywords[:5]
    if job_titles:
        search_keywords = job_titles[:2] + search_keywords

    def fetch_adzuna_all():
        """Fetch from multiple regions for broader coverage."""
        jobs = []
        for country in ("gb", "us", "de"):
            jobs.extend(fetch_jobs_from_adzuna(search_keywords, country=country, per_page=5))
        return jobs[:15]

    def fetch_jsearch():
        return fetch_jobs_from_jsearch(search_keywords, location="", per_page=10)

    def fetch_arbeitnow():
        return fetch_jobs_from_arbeitnow(search_keywords, location="", per_page=10)

    job_lists = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        sources = [
            ("adzuna", fetch_adzuna_all),
            ("jsearch", fetch_jsearch),
            ("arbeitnow", fetch_arbeitnow),
        ]
        future_to_name = {executor.submit(fn): name for name, fn in sources}
        for fut in concurrent.futures.as_completed(future_to_name):
            name = future_to_name[fut]
            try:
                result = fut.result()
                logging.info("jobs_recommend: %s returned %d jobs", name, len(result) if isinstance(result, list) else -1)
                job_lists.append(result if isinstance(result, list) else [])
            except Exception as e:
                logging.warning("jobs_recommend: %s fetch failed: %s", name, e)
                job_lists.append([])

    jobs = merge_and_deduplicate(job_lists, max_jobs=25)
    jobs = rank_jobs_by_relevance(jobs, resume_text, max_to_rank=15)
    sources = list(dict.fromkeys(j.get("source", "") for j in jobs if j.get("source")))
    return jsonify({"jobs": jobs, "sources": sources})

IMPROVE_SYSTEM = """You are an expert resume writer. Given a resume, job description, and optional feedback, produce an IMPROVED version as plain text. Do NOT invent facts. Use section headers in ALL CAPS on their own line: SUMMARY, EXPERIENCE, EDUCATION, SKILLS. Output only the improved resume text."""

model_improve = genai.GenerativeModel(
    model_name='gemini-2.5-flash',
    system_instruction=IMPROVE_SYSTEM,
)

CV_JSON_SYSTEM = (
    "You are an expert resume writer. You ONLY output valid JSON. "
    "No markdown, no code fences, no explanation, no trailing commas. "
    "Given a resume and job description, return a single JSON object with these exact keys: "
    "name, title, contact (object with email/phone/location/website), summary, "
    "skills (array of strings), experience (array of objects with role/company/location/dates/bullets), "
    "education (array of objects with degree/school/year), "
    "projects (array of objects with title/description), hobbies (array of strings). "
    "Improve the content to match the job description. Do NOT invent facts."
)

model_cv_json = genai.GenerativeModel(
    model_name='gemini-2.5-flash',
    system_instruction=CV_JSON_SYSTEM,
)

def build_docx_from_text(text):
    """Build a simple DOCX from structured resume text (headers in ALL CAPS, then content)."""
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'
    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.isupper() and len(line) > 1 and line.strip():
            # Section heading
            p = doc.add_paragraph()
            p.add_run(line.strip()).bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        # Content
        content = []
        while i < len(lines) and (not lines[i].isupper() or len(lines[i]) <= 1):
            content.append(lines[i].strip())
            i += 1
        block = "\n".join(content).strip()
        if block:
            for part in block.split("\n"):
                if part.strip():
                    doc.add_paragraph(part.strip())
    if not doc.paragraphs:
        doc.add_paragraph(text[:5000])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf_from_text(text):
    """Build a professional PDF from structured resume text (headers in ALL CAPS, then content)."""
    from fpdf import FPDF

    class CVPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=20)

        def header(self):
            pass

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            w = self.w - self.l_margin - self.r_margin
            if w > 0:
                self.cell(w, 10, "CV generated by CVSmart", align="C")

    pdf = CVPDF()
    pdf.set_margins(25, 20, 25)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(0, 0, 0)

    # Explicit width so multi_cell never gets 0 (avoids "Not enough horizontal space")
    usable_w = max(1, (pdf.w - pdf.l_margin - pdf.r_margin))

    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.isupper() and len(line) > 1 and line.strip():
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(4)
            pdf.multi_cell(usable_w, 8, line.strip())
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 11)
            i += 1
            continue
        content = []
        while i < len(lines) and (not lines[i].isupper() or len(lines[i]) <= 1):
            content.append(lines[i].strip())
            i += 1
        block = "\n".join(content).strip()
        if block:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(40, 40, 40)
            for part in block.split("\n"):
                s = part.strip()
                if s:
                    pdf.multi_cell(usable_w, 6, s)
            pdf.ln(2)
    if pdf.page_no() == 1 and not lines:
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(usable_w, 6, text[:5000])
    buffer = io.BytesIO()
    buffer.write(pdf.output())
    buffer.seek(0)
    return buffer


def _latin1(text_value):
    """Safely encode text for FPDF's built-in Helvetica (latin-1 only)."""
    if not text_value:
        return ""
    s = str(text_value)
    return s.encode("latin-1", "replace").decode("latin-1")


def build_cv_pdf_from_sections(sections, template_id="classic"):
    """Build a professional CV PDF from structured sections dict."""
    from fpdf import FPDF

    class CVTemplatePDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=18)
        def header(self):
            pass
        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(160, 160, 160)
            w = self.w - self.l_margin - self.r_margin
            if w > 0:
                self.cell(w, 8, "Generated by CVSmart", align="C")

    pdf = CVTemplatePDF()
    pdf.set_margins(22, 18, 22)
    pdf.add_page()
    W = max(1, pdf.w - pdf.l_margin - pdf.r_margin)

    name = _latin1(sections.get("name") or "")
    title = _latin1(sections.get("title") or "")
    contact = sections.get("contact") or {}
    if not isinstance(contact, dict):
        contact = {}
    contact_parts = []
    for k in ("email", "phone", "location", "website"):
        v = _latin1(contact.get(k) or "").strip()
        if v:
            contact_parts.append(v)
    contact_line = "  |  ".join(contact_parts)
    summary = _latin1(sections.get("summary") or "")
    skills = sections.get("skills") or []
    experience = sections.get("experience") or []
    education = sections.get("education") or []
    projects = sections.get("projects") or []
    hobbies = sections.get("hobbies") or []

    # Colors per template
    if template_id == "modern":
        hdr_bg = (52, 58, 64)
        hdr_fg = (255, 255, 255)
        accent = (0, 123, 255)
        body_fg = (33, 37, 41)
        line_col = (0, 123, 255)
    elif template_id == "minimal":
        hdr_bg = None
        hdr_fg = (0, 0, 0)
        accent = (80, 80, 80)
        body_fg = (50, 50, 50)
        line_col = (200, 200, 200)
    else:
        hdr_bg = (245, 245, 245)
        hdr_fg = (0, 0, 0)
        accent = (30, 80, 160)
        body_fg = (40, 40, 40)
        line_col = (30, 80, 160)

    def _hr():
        pdf.set_draw_color(*line_col)
        y = pdf.get_y()
        pdf.line(pdf.l_margin, y, pdf.l_margin + W, y)
        pdf.ln(3)

    # ── NAME / TITLE / CONTACT ──
    if template_id == "modern" and hdr_bg:
        pdf.set_fill_color(*hdr_bg)
        pdf.set_text_color(*hdr_fg)
        pdf.set_font("Helvetica", "B", 22)
        pdf.cell(W, 13, name or "Candidate", ln=1, fill=True)
        if title:
            pdf.set_font("Helvetica", "", 12)
            pdf.cell(W, 8, title, ln=1, fill=True)
        if contact_line:
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(W, 7, contact_line, ln=1, fill=True)
        pdf.ln(6)
    elif template_id == "minimal":
        pdf.set_text_color(*hdr_fg)
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(W, 11, name or "Candidate", ln=1)
        if title:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(*accent)
            pdf.cell(W, 6, title, ln=1)
        if contact_line:
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(120, 120, 120)
            pdf.cell(W, 5, contact_line, ln=1)
        pdf.ln(2)
        _hr()
    else:
        if hdr_bg:
            pdf.set_fill_color(*hdr_bg)
        pdf.set_text_color(*hdr_fg)
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(W, 12, name or "Candidate", ln=1, fill=bool(hdr_bg))
        if title:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(*accent)
            pdf.cell(W, 7, title, ln=1, fill=bool(hdr_bg))
        if contact_line:
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(W, 6, contact_line, ln=1, fill=bool(hdr_bg))
        pdf.ln(4)
        _hr()

    def section_heading(label):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*accent)
        pdf.cell(W, 7, label.upper(), ln=1)
        pdf.set_draw_color(*line_col)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + W, pdf.get_y())
        pdf.ln(2)

    def body_text(txt, bold=False):
        pdf.set_font("Helvetica", "B" if bold else "", 10)
        pdf.set_text_color(*body_fg)
        pdf.multi_cell(W, 5, _latin1(txt))

    def bullet(txt):
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*body_fg)
        pdf.cell(5, 5, "-")
        x = pdf.get_x()
        pdf.multi_cell(W - 5, 5, _latin1(txt))

    # ── SUMMARY ──
    if summary:
        section_heading("Professional Summary")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*body_fg)
        pdf.multi_cell(W, 5, summary)
        pdf.ln(1)

    # ── EXPERIENCE ──
    if isinstance(experience, list) and experience:
        section_heading("Experience")
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            role = _latin1(exp.get("role") or "")
            company = _latin1(exp.get("company") or "")
            dates = _latin1(exp.get("dates") or "")
            loc = _latin1(exp.get("location") or "")
            line1 = role
            if company:
                line1 = f"{role}  -  {company}" if role else company
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*body_fg)
            pdf.cell(W * 0.7, 5, line1)
            if dates:
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(120, 120, 120)
                pdf.cell(W * 0.3, 5, dates, align="R")
            pdf.ln(5)
            if loc:
                pdf.set_font("Helvetica", "I", 8)
                pdf.set_text_color(130, 130, 130)
                pdf.cell(W, 4, loc, ln=1)
            bullets = exp.get("bullets") or []
            if isinstance(bullets, list):
                for b in bullets:
                    s = str(b or "").strip()
                    if s:
                        bullet(s)
            pdf.ln(2)

    # ── EDUCATION ──
    if isinstance(education, list) and education:
        section_heading("Education")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            deg = _latin1(edu.get("degree") or "")
            sch = _latin1(edu.get("school") or "")
            yr = _latin1(edu.get("year") or "")
            line1 = deg
            if sch:
                line1 = f"{deg}  -  {sch}" if deg else sch
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*body_fg)
            pdf.cell(W * 0.7, 5, line1)
            if yr:
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(120, 120, 120)
                pdf.cell(W * 0.3, 5, yr, align="R")
            pdf.ln(6)

    # ── SKILLS ──
    if isinstance(skills, list) and skills:
        section_heading("Skills")
        safe_skills = [_latin1(str(s).strip()) for s in skills if str(s).strip()]
        chunks = [safe_skills[i:i+5] for i in range(0, len(safe_skills), 5)]
        for chunk in chunks:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(*body_fg)
            pdf.multi_cell(W, 5, "  |  ".join(chunk))
        pdf.ln(1)

    # ── PROJECTS ──
    if isinstance(projects, list) and projects:
        section_heading("Projects")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            t = _latin1(proj.get("title") or "")
            d = _latin1(proj.get("description") or "")
            if t:
                body_text(t, bold=True)
            if d:
                body_text(d)
            pdf.ln(2)

    # ── HOBBIES ──
    if isinstance(hobbies, list) and hobbies:
        section_heading("Interests")
        safe_hobbies = [_latin1(str(h).strip()) for h in hobbies if str(h).strip()]
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*body_fg)
        pdf.multi_cell(W, 5, "  |  ".join(safe_hobbies))

    buf = io.BytesIO()
    buf.write(pdf.output())
    buf.seek(0)
    return buf


@app.route('/cv/improve', methods=['POST'])
@limiter.limit("10 per minute")
def cv_improve():
    """Generate improved CV as DOCX from resume + job description + optional analysis."""
    if 'resume' not in request.files or not request.files['resume']:
        return jsonify({'error': 'No resume file uploaded'}), 400
    job_description = request.form.get('jobDescription', '').strip()
    if not job_description or len(job_description) < 10:
        return jsonify({'error': 'Job description required'}), 400
    analysis_text = request.form.get('analysis', '').strip()

    file = request.files['resume']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    try:
        if filename.endswith('.pdf'):
            resume_text = extract_text_from_pdf(file_path)
        else:
            resume_text = extract_text_from_docx(file_path)
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

    if not resume_text.strip():
        return jsonify({'error': 'Could not extract text from resume'}), 400

    prompt = f"""Resume:\n{resume_text[:6000]}\n\nJob description:\n{job_description[:3000]}\n\n"""
    if analysis_text:
        prompt += f"Feedback to apply:\n{analysis_text[:3000]}\n\n"
    prompt += "Output the improved resume as plain text with section headers in ALL CAPS."

    try:
        response = model_improve.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(temperature=0.3, max_output_tokens=4000),
        )
        improved_text = (response.text or "").strip()
        if not improved_text:
            return jsonify({'error': 'Failed to generate improved resume'}), 500
        want_pdf = request.form.get('format', '').strip().lower() == 'pdf'
        if want_pdf:
            buffer = build_pdf_from_text(improved_text)
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='improved_cv.pdf',
            )
        buffer = build_docx_from_text(improved_text)
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='improved_cv.docx',
        )
    except Exception as e:
        logging.exception("cv_improve error: %s", e)
        return jsonify({'error': str(e)}), 500


@app.route('/cv/templates', methods=['POST'])
@limiter.limit("10 per minute")
def cv_templates():
    """Return structured CV sections JSON generated by Gemini."""
    if 'resume' not in request.files or not request.files['resume']:
        return jsonify({'error': 'No resume file uploaded'}), 400
    job_description = (request.form.get('jobDescription') or '').strip()
    if not job_description or len(job_description) < 10:
        return jsonify({'error': 'Job description required'}), 400
    analysis_text = (request.form.get('analysis') or '').strip()

    file = request.files['resume']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    try:
        file.save(file_path)
        if filename.endswith('.pdf'):
            resume_text = extract_text_from_pdf(file_path)
        else:
            resume_text = extract_text_from_docx(file_path)
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

    if not resume_text.strip():
        return jsonify({'error': 'Could not extract text from resume'}), 400

    try:
        sections = generate_structured_cv_sections(resume_text, job_description, analysis_text)
    except Exception as e:
        msg = str(e)
        logging.exception("cv/templates error: %s", e)
        if "429" in msg or "ResourceExhausted" in msg:
            return jsonify({'error': 'Gemini quota exceeded – please try again later.'}), 503
        # Surface the underlying error to the client to make debugging easier.
        return jsonify({'error': msg or 'Failed to generate CV structure'}), 500

    return jsonify({'sections': sections})

def build_cv_docx(sections):
    """Build DOCX from CV builder sections dict."""
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    def add_heading(text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def add_para(text):
        if text and str(text).strip():
            doc.add_paragraph(str(text).strip())

    summary = sections.get("summary") or ""
    if summary:
        add_heading("SUMMARY")
        add_para(summary)

    experience = sections.get("experience") or []
    if experience:
        add_heading("EXPERIENCE")
        for exp in experience:
            if isinstance(exp, dict):
                role = exp.get("role") or ""
                company = exp.get("company") or ""
                dates = exp.get("dates") or ""
                bullets = exp.get("bullets") or []
                if role or company:
                    add_para(f"{role} at {company}" + (f" — {dates}" if dates else ""))
                for b in bullets if isinstance(bullets, list) else [bullets]:
                    if b and str(b).strip():
                        doc.add_paragraph(str(b).strip(), style='List Bullet')
            elif isinstance(exp, str) and exp.strip():
                add_para(exp)

    education = sections.get("education") or []
    if education:
        add_heading("EDUCATION")
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree") or ""
                school = edu.get("school") or ""
                year = edu.get("year") or ""
                add_para(f"{degree} — {school}" + (f", {year}" if year else ""))
            elif isinstance(edu, str) and edu.strip():
                add_para(edu)

    skills = sections.get("skills") or []
    if skills:
        add_heading("SKILLS")
        if isinstance(skills, list):
            add_para(", ".join(str(s).strip() for s in skills if s and str(s).strip()))
        else:
            add_para(str(skills))

    projects = sections.get("projects") or []
    if projects:
        add_heading("PROJECTS")
        for proj in projects:
            if isinstance(proj, dict):
                title = proj.get("title") or ""
                desc = proj.get("description") or ""
                add_para(title)
                if desc:
                    add_para(desc)
            elif isinstance(proj, str) and proj.strip():
                add_para(proj)

    if not doc.paragraphs:
        doc.add_paragraph("No content added yet.")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/cv/build', methods=['POST'])
@limiter.limit("20 per minute")
def cv_build():
    """Build CV DOCX from JSON (templateId + sections)."""
    try:
        data = request.get_json() or {}
    except Exception:
        return jsonify({"error": "Invalid JSON"}), 400
    template_id = data.get("templateId", "classic")
    sections = data.get("sections") or {}
    if not isinstance(sections, dict):
        return jsonify({"error": "sections must be an object"}), 400
    try:
        buffer = build_cv_docx(sections)
        from flask import send_file
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='my_cv.docx',
        )
    except Exception as e:
        logging.exception("cv/build error: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route('/cv/build/pdf', methods=['POST'])
@limiter.limit("10 per minute")
def cv_build_pdf():
    """Build CV PDF from JSON (templateId + sections)."""
    try:
        data = request.get_json() or {}
    except Exception:
        return jsonify({"error": "Invalid JSON"}), 400
    template_id = str(data.get("templateId") or "classic").strip().lower()
    sections = data.get("sections") or {}
    if not isinstance(sections, dict):
        return jsonify({"error": "sections must be an object"}), 400
    if template_id not in {"classic", "modern", "minimal"}:
        template_id = "classic"
    try:
        buffer = build_cv_pdf_from_sections(sections, template_id)
        return send_file(
            buffer,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="improved_cv.pdf",
        )
    except Exception as e:
        logging.exception("cv/build/pdf error: %s", e)
        return jsonify({"error": str(e)}), 500

ASSESS_PROMPT = """Based on the job description below, generate exactly 6 multiple-choice questions that candidates are likely to be asked in interviews for this role.
Include a mix of: role-specific technical or situational questions, behavioral questions relevant to the responsibilities, and questions about skills/experience mentioned in the JD.
For each question provide:
- question: the question text
- options: an object with keys A, B, C, D and string values (the answer choices)
- correct: exactly one of "A", "B", "C", "D"
- explanation: a short explanation or tip for answering (1-2 sentences)

Output ONLY a single JSON object. No markdown, no code fences, no text before or after. Valid JSON only:
{"questions": [{"question": "Question text?", "options": {"A": "Choice A", "B": "Choice B", "C": "Choice C", "D": "Choice D"}, "correct": "A", "explanation": "Short explanation."}]}"""


def _parse_assess_response(text):
    """Parse Gemini response into questions list."""
    text = (text or "").strip()
    if not text:
        return []
    # Strip markdown code block if present
    if "```" in text:
        parts = re.split(r"```(?:\w*)\s*\n?", text, maxsplit=1)
        if len(parts) > 1:
            text = parts[1]
        text = re.sub(r"\s*```\s*$", "", text)
    text = text.strip()
    if not text:
        return []
    # Extract JSON object
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return []
    text = text[start : end + 1]
    # Fix common JSON issues
    text_fixed = re.sub(r",\s*([}\]])", r"\1", text)
    try:
        data = json.loads(text_fixed)
    except json.JSONDecodeError:
        return []
    questions = data.get("questions") or []
    if not isinstance(questions, list):
        return []
    return questions


def _parse_assess_questions(raw_questions):
    """Validate and normalize questions to expected shape."""
    out = []
    for q in raw_questions[:8]:
        if not isinstance(q, dict):
            continue
        opts = q.get("options") or {}
        # Accept options with keys A,B,C,D or 0,1,2,3 or first 4 keys
        option_keys = list(opts.keys()) if isinstance(opts, dict) else []
        if len(option_keys) < 2:
            continue
        # Map to A,B,C,D
        key_map = {}
        for i, k in enumerate(option_keys[:4]):
            key_map[k] = ("A", "B", "C", "D")[i]
        normalized = {}
        for k, v in opts.items():
            if k in key_map:
                normalized[key_map[k]] = str(v) if v else ""
        for letter in ("A", "B", "C", "D"):
            if letter not in normalized:
                normalized[letter] = ""
        correct = str(q.get("correct", "")).upper()
        if correct in key_map:
            correct = key_map[correct]
        elif correct not in ("A", "B", "C", "D"):
            correct = "A" if normalized.get("A") else list(normalized.keys())[0]
        out.append({
            "question": str(q.get("question", ""))[:500],
            "options": {k: str(normalized.get(k, ""))[:200] for k in ("A", "B", "C", "D")},
            "correct": correct,
            "explanation": str(q.get("explanation", ""))[:300],
        })
    return out


@app.route('/assess/generate', methods=['POST'])
@limiter.limit("10 per minute")
def assess_generate():
    """Generate interview preparation MCQ questions from job description (and optional resume)."""
    job_description = request.form.get("jobDescription", "").strip()
    if not job_description or len(job_description) < 20:
        return jsonify({"error": "Job description required (min 20 chars)"}), 400

    resume_text = request.form.get("resumeText", "").strip()
    file = request.files.get("resume") if "resume" in request.files else None
    if file and file.filename and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        try:
            file.save(file_path)
            resume_text = extract_text_from_pdf(file_path) if filename.endswith(".pdf") else extract_text_from_docx(file_path)
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)

    # Instructions first, then data (better for model)
    prompt = ASSESS_PROMPT + "\n\n---\n\nJob description:\n" + job_description[:4000]
    if resume_text:
        prompt += "\n\nCandidate resume (for context):\n" + resume_text[:2000]

    for attempt in range(2):
        try:
            config_kw = {"temperature": 0.3 if attempt == 0 else 0.2, "max_output_tokens": 4000}
            if attempt == 0:
                config_kw["response_mime_type"] = "application/json"
            config = genai.GenerationConfig(**config_kw)
            response = model.generate_content(prompt, generation_config=config)
            text = (response.text or "").strip()
            if not text:
                logging.warning("assess/generate empty response from model")
                continue

            raw = _parse_assess_response(text)
            out = _parse_assess_questions(raw)

            # Also try strict format if lenient parsing got nothing
            if not out and raw:
                for q in raw[:8]:
                    if not isinstance(q, dict):
                        continue
                    opts = q.get("options") or {}
                    if not all(k in opts for k in ("A", "B", "C", "D")):
                        continue
                    correct = q.get("correct")
                    if correct not in ("A", "B", "C", "D"):
                        continue
                    out.append({
                        "question": str(q.get("question", "")),
                        "options": {k: str(opts.get(k, "")) for k in ("A", "B", "C", "D")},
                        "correct": correct,
                        "explanation": str(q.get("explanation", "")),
                    })

            if out:
                return jsonify({"questions": out})

        except json.JSONDecodeError as e:
            logging.warning("assess/generate JSON parse error (attempt %d): %s", attempt + 1, e)
        except Exception as e:
            logging.exception("assess/generate error (attempt %d): %s", attempt + 1, e)
            if attempt == 0:
                # Retry once with a simpler config
                continue
            msg = str(e)
            lower = msg.lower()
            if "429" in msg or "quota" in lower or "rate limit" in lower:
                return jsonify({
                    "error": "AI quota has been exceeded for interview prep. Please wait a minute and try again.",
                    "questions": [],
                }), 503
            return jsonify({
                "error": "Interview questions could not be generated. Please try again later.",
                "questions": [],
            }), 500

    return jsonify({"error": "Failed to generate questions", "questions": []}), 200

if __name__ == '__main__':
    app.run(debug=True)
